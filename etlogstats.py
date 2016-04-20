#!/usr/bin/env python

# Import libraries
import glob 
import os
from geolite2 import geolite2
from operator import itemgetter
import collections
from collections import OrderedDict
import csv
from docx import Document
import pprint
import hashlib

# Main class
class ETLogStats():
	'Generate stats from the online and offline web logs for etekkatho'
	
	def __init__(self):
		#self.onlineStats()
		self.offlineStats()
		print('Done.')

	def onlineStats(self):
		print('Analysing online web logs...')

		linesProcessed = 0
		path = './online/*'   
		files = glob.glob(path)  
		startDate = '02 March 2015'
		endDate = '04 March 2016'
		pdfCount = 0
		pdfCountMyanmar = 0
		pdfMyanmarPercentage = 0
		ipList = []
		monthTotals = {}
		myanmarMonthTotals = {}
		pageViews = 0
		
		# Loop through log files
		for file in files:     
			#print(os.path.basename(path))
		
			# Open the first file
			with open(file) as fp:			
				for line in fp:
					linesProcessed = linesProcessed+1
					print('Processing log file row: '+str(linesProcessed))
					
					if '200' in line and not any(x in line for x in ('.zip', '.pdf', '.css', '.js', '.png', '.jpg', '.jpeg', '.gif', '.mp4', '.ico', '.txt', '.php', 'bot', 'spider')):
						pageViews = pageViews+1
					
					lineData = self.apache2_logrow(line)
					#print('.', end='')
					
					# If requested file is a PDF add to the count
					if "pdf" in lineData[4] or "zip" in lineData[4]:
						#print(lineData[4])
						pdfCount = pdfCount+1
						#print('*', end='')
						
						# Add IP to list
						ipList.append(lineData[0])
						
						logDate = lineData[3].split('/')
						logMonth = logDate[1]
						logYear = logDate[2].split(':')
						logYear = logYear[0]
						
						# All month totals
						if logYear not in monthTotals:
							monthTotals[logYear] = {}
						
						if logMonth in monthTotals[logYear]:
							monthTotals[logYear][logMonth] = monthTotals[logYear][logMonth]+1
						else:
							monthTotals[logYear][logMonth] = 1
							
						# Myanmar month totals
						reader = geolite2.reader()
						geo = reader.get(lineData[0])
						if 'country' in geo:
							country = geo['country']['iso_code']
							if 'MM' in country:
							
								if logYear not in myanmarMonthTotals:
									myanmarMonthTotals[logYear] = {}
						
								if logMonth in myanmarMonthTotals[logYear]:
									myanmarMonthTotals[logYear][logMonth] = myanmarMonthTotals[logYear][logMonth]+1
								else:
									myanmarMonthTotals[logYear][logMonth] = 1
						
		
		# Work out the Myanmar IPs
		ipCountries = {}
		
		for ip in ipList:
			#print(ip)
		
			if ip is not '::1':
				reader = geolite2.reader()
				geo = reader.get(ip)
			
				if 'country' in geo:
					country = geo['country']['iso_code']
					countryName = geo['country']['names']['en']
					#print(country)
					#quit()
					
					#print(type(country))
					
					if 'MM' in country:
						#print(country)
						pdfCountMyanmar = pdfCountMyanmar+1
					
					if countryName in ipCountries:
						ipCountries[countryName] = ipCountries[countryName]+1
					else:
						ipCountries[countryName] = 1
		
		geolite2.close()

		# Work out the Myanmar downloads as a percentage
		if pdfCountMyanmar > 0:
			pdfCountMyanmar1percent = pdfCount/100
			pdfMyanmarPercentage = pdfCountMyanmar/pdfCountMyanmar1percent
		
		print('\nWeb log analysis complete:')
		print('\t Log lines processed: {}'.format(linesProcessed))
		print('\t {} to {}'.format(startDate, endDate))
		print('\t Total page views: {}'.format(pageViews))
		print('\t Total PDF and ZIP downloads {}'.format(pdfCount))
		print('\t Total PDF and ZIP downloads from Myanmar {}'.format(pdfCountMyanmar))
		print('\t Myanmar percentage {}%'.format(round(pdfMyanmarPercentage, 2)))
		print('\t Downloads by country: ')
		
		sorted_ipCountries = OrderedDict(sorted(ipCountries.items(), key=itemgetter(1), reverse=True))
		for x in sorted_ipCountries:
			print('\t\t - {} : {}'.format(x, ipCountries[x]))
	
		pp = pprint.PrettyPrinter(indent=3)
		pp.pprint(monthTotals)
		
		for year in monthTotals:		
			for month in monthTotals[year]:
				print(year+', '+month+', '+str(monthTotals[year][month]))
		
		
		# Create CSV 
		with open('onlinestats.csv', 'w', newline='') as csvfile: 
			onlineCSV = csv.writer(csvfile, delimiter=',', quotechar='"', quoting=csv.QUOTE_MINIMAL)
			
			onlineCSV.writerow(['Time period ', startDate+' to '+endDate])
			onlineCSV.writerow(['Total downloads', pdfCount])
			
			pmp = str(round(pdfMyanmarPercentage, 2))
			onlineCSV.writerow(['Total downloads from Myanmar', str(pdfCountMyanmar)+' ('+str(pmp+'%')+')'])
			
			onlineCSV.writerow(['Total downloads by month:'])
			onlineCSV.writerow(['Year', 'Month', 'Downloads'])
			
			for year in monthTotals:		
				for month in monthTotals[year]:
					onlineCSV.writerow([year, month, monthTotals[year][month]])
					
			onlineCSV.writerow(['Myanmar downloads by month:'])
			onlineCSV.writerow(['Year', 'Month', 'Downloads'])
			
			for year in myanmarMonthTotals:		
				for month in myanmarMonthTotals[year]:
					onlineCSV.writerow([year, month, myanmarMonthTotals[year][month]])
					
			onlineCSV.writerow(['Country', 'Downloads'])
			sorted_ipCountries = OrderedDict(sorted(ipCountries.items(), key=itemgetter(1), reverse=True))
			for x in sorted_ipCountries:
				onlineCSV.writerow([x, ipCountries[x]])
					
			
		
		# Create document
		'''
		document = Document()
		document.add_heading('eTekktho.org Online Usage', 0)
		
		document.add_heading(startDate+' to '+endDate, level=1)
		
		document.add_paragraph('Total resource downloads: '+str(pdfCount))
		pmp = str(round(pdfMyanmarPercentage, 2))
		document.add_paragraph('Total resource downloads from Myanmar: '+str(pdfCountMyanmar)+' ('+str(pmp+'%')+')')
		
		
		document.add_heading('Total downloads by month', level=2)
		
		table = document.add_table(rows=1, cols=3)
		hdr_cells = table.rows[0].cells
		hdr_cells[0].text = 'Year'
		hdr_cells[1].text = 'Month'
		hdr_cells[2].text = 'Downloads'
		for year in monthTotals:
			for month in monthTotals[year]:
				row_cells = table.add_row().cells
				row_cells[0].text = str(year)
				row_cells[1].text = str(month)
				row_cells[2].text = str(monthTotals[year][month])
			
		document.save('et_onlinestats.docx')
		'''
		
		
	def offlineStats(self):
		print('Analysing offline web logs...')
		
		# Vars
		rootdir = './offline/'
		hashs = []
		logsChecked = 0
		logsDuplicate = 0
		logsUnique = 0
		downloadTotal = 0
		duplicateFiles = {}
		monthTotals = {}
		pageView = 0
		logLineCount = 0
		totalLinesRead = 0
		videoDownloads = 0
		
		# Loop through files
		for subdir, dirs, files in os.walk(rootdir):
			for file in files:
				#print(os.path.join(subdir, file))
				
				# Only check log files
				if file.endswith('.log'):
					#print(os.path.join(subdir, file))
					logsChecked = logsChecked+1
					
					# Hash file and filename to check for duplicates
					filepath = os.path.join(subdir, file)
					hash = self.hash_file(filepath)
					
					#print(hash)
					
					if hash not in hashs:
						duplicateFiles[hash] = {}
						duplicateFiles[hash][0] = filepath
						
						logsUnique = logsUnique+1
						
						# Add hash
						hashs.append(hash)
						
						# Read log file
						with open(filepath) as fp:
							for line in fp:
								totalLinesRead = totalLinesRead+1
								
								# Check for a PDF or ZIP
								if line.startswith('#'):
									# Header
									pass
								else:
									# Log line
									logLineCount = logLineCount+1
									
									if "pdf" in line or "zip" in line or "mp4" in line: 
										downloadTotal = downloadTotal+1
										
										# Total month totals
										logDate = line.split(' ')
										logDate = logDate[0].split('-')
										logMonth = logDate[1]
										logYear = logDate[0]
										
										if logYear not in monthTotals:
											monthTotals[logYear] = {}
						
										if logMonth in monthTotals[logYear]:
											monthTotals[logYear][logMonth] = monthTotals[logYear][logMonth]+1
										else:
											monthTotals[logYear][logMonth] = 1
											
									if '200' in line and not any(x in line for x in ('.zip', '.pdf', '.css', '.js', '.png', '.jpg', '.jpeg', '.gif', '.mp4', '.ico')) and '.html' in line:
									#if '200' in line and '.html' in line:
										pageView = pageView+1
									
									if 'mp4' in line:
										videoDownloads = videoDownloads+1
									
										
					
					else:
						#print('Duplicate hash for: '+filepath)
						logsDuplicate = logsDuplicate+1
						
						duplicateFiles[hash][1] = filepath
		
		'''
		for x in duplicateFiles:	
			if 1 in duplicateFiles[x]:
				print('Duplicate log file:')
				print(duplicateFiles[x][0])
				print(duplicateFiles[x][1])
		'''
		print('Logs checked: '+str(logsChecked))
		print('Logs duplicate: '+str(logsDuplicate))
		print('Logs unique: '+str(logsUnique))
		
		print('Total lines read: '+str(totalLinesRead))
		print('Non header lines read: '+str(logLineCount))
		
		print('Page views total: '+str(pageView))
		
		print('Download video: '+str(videoDownloads))

		print('Download total: '+str(downloadTotal))
		
		print('Downloads by month:')
		#pp = pprint.PrettyPrinter(indent=3)
		#pp.pprint(monthTotals)
		
		od = collections.OrderedDict(sorted(monthTotals.items()))
		
		for year, month in od.items():	
			odm = collections.OrderedDict(sorted(month.items()))
			
			for mon, value in odm.items():
				print(year+', '+mon+', '+str(value))
		
		# Create CSV 
		with open('offlinestats.csv', 'w', newline='') as csvfile: 
			onlineCSV = csv.writer(csvfile, delimiter=',', quotechar='"', quoting=csv.QUOTE_MINIMAL)
			
			onlineCSV.writerow(['Time period ', '2013 to 2016'])
			onlineCSV.writerow(['Total downloads (PDF, ZIP, MP4)', downloadTotal])
			onlineCSV.writerow(['Total page views', pageView])
			
			onlineCSV.writerow(['Total downloads by month:'])
			onlineCSV.writerow(['Year', 'Month', 'Downloads'])
					
			od = collections.OrderedDict(sorted(monthTotals.items()))
			monthNames = {
				1: 'Jan',
				2: 'Feb',
				3: 'Mar',
				4: 'Apr',
				5: 'May',
				6: 'Jun',
				7: 'Jul',
				8: 'Aug',
				9: 'Sep',
				10: 'Oct',
				11: 'Nov',
				12: 'Dec'
			}
			
			for year, month in od.items():	
				odm = collections.OrderedDict(sorted(month.items()))
			
				for mon, value in odm.items():
					onlineCSV.writerow([year,monthNames[int(mon)], value])
					
			
	
		
	def apache2_logrow(self, s):
		''' Fast split on Apache2 log lines

		http://httpd.apache.org/docs/trunk/logs.html
		'''
		row = [ ]
		qe = qp = None # quote end character (qe) and quote parts (qp)
		for s in s.replace('\r','').replace('\n','').split(' '):
			if qp:
				qp.append(s)
			elif '' == s: # blanks
				row.append('')
			elif '"' == s[0]: # begin " quote "
				qp = [ s ]
				qe = '"'
			elif '[' == s[0]: # begin [ quote ]
				qp = [ s ]
				qe = ']'
			else:
				row.append(s)

			l = len(s)
			if l and qe == s[-1]: # end quote
				if l == 1 or s[-2] != '\\': # don't end on escaped quotes
					row.append(' '.join(qp)[1:-1].replace('\\'+qe, qe))
					qp = qe = None
		return row
		
	def hash_file(self, filename):
	   """"This function returns the SHA-1 hash
	   of the file passed into it"""

	   # make a hash object
	   h = hashlib.sha1()

	   # open file for reading in binary mode
	   with open(filename,'rb') as file:

		   # loop till the end of the file
		   chunk = 0
		   while chunk != b'':
			   # read only 1024 bytes at a time
			   chunk = file.read(1024)
			   h.update(chunk)

	   # return the hex representation of digest
	   return h.hexdigest()

ETLogStats()
